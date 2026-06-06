from __future__ import annotations

from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import subprocess
import zipfile
import xml.etree.ElementTree as ET

import fitz  # PyMuPDF

try:
    from ..models import IssueItem, ParseResult
    from .markdown_annotator import build_annotated_markdown, markdown_to_pdf, write_markdown_bundle
    from .office import find_soffice
    from .pdf_review_annotator import annotate_pdf_from_issues
    from .temp_paths import mkdtemp
except ImportError:  # pragma: no cover
    from models import IssueItem, ParseResult
    from services.markdown_annotator import build_annotated_markdown, markdown_to_pdf, write_markdown_bundle
    from services.office import find_soffice
    from services.pdf_review_annotator import annotate_pdf_from_issues
    from services.temp_paths import mkdtemp


def _build_comment_text(issue: IssueItem) -> str:
    return (
        f"[Category] {issue.category}\n"
        f"[Severity] {issue.severity}\n"
        f"[Original Text] {issue.original_text or issue.snippet or issue.review_input_text}\n"
        f"[Reason] {issue.reason}\n"
        f"[Suggestion] {issue.suggestion}\n"
        f"[Confidence] {issue.confidence:.2f}"
    )


def _apply_docx_annotations(parse_result: ParseResult, issues: List[IssueItem], output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_COLOR_INDEX
    except Exception:
        return _apply_docx_annotations_openxml(parse_result, issues, output_path)

    doc = Document(str(parse_result.normalized_path))

    issue_map: Dict[str, List[IssueItem]] = defaultdict(list)
    para_anchor = {p.para_id: p.anchor for p in parse_result.paragraphs}
    for issue in issues:
        issue_map[issue.para_id].append(issue)

    for para_id, anchor in para_anchor.items():
        if para_id not in issue_map:
            continue

        if anchor.get("kind") in {"paragraph", "word_main_paragraph"}:
            p_idx = int(anchor["paragraph_index"])
            paragraph = doc.paragraphs[p_idx]
        elif anchor.get("kind") in {"table_cell", "word_main_table_cell"}:
            table = doc.tables[int(anchor["table_index"])]
            row = table.rows[int(anchor["row_index"])]
            cell = row.cells[int(anchor["col_index"])]
            paragraph = cell.paragraphs[int(anchor["paragraph_index"])]
        else:
            continue

        if not paragraph.runs:
            paragraph.add_run(" ")

        for issue in issue_map[para_id]:
            comment_text = _build_comment_text(issue)
            marker_run = paragraph.add_run(f" [{issue.issue_no}]")
            marker_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            try:
                doc.add_comment(runs=paragraph.runs, text=comment_text, author="DocReview", initials="DR")
            except Exception:
                paragraph.add_run(f" [REVIEW NOTE] {comment_text}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _apply_docx_annotations_openxml(parse_result: ParseResult, issues: List[IssueItem], output_path: Path) -> Path:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    ET.register_namespace("w", w_ns)

    issue_map: Dict[str, List[IssueItem]] = defaultdict(list)
    para_anchor = {p.para_id: p.anchor for p in parse_result.paragraphs}
    for issue in issues:
        issue_map[issue.para_id].append(issue)

    with zipfile.ZipFile(str(parse_result.normalized_path), "r") as zf:
        archive_data = {name: zf.read(name) for name in zf.namelist()}

    if "word/document.xml" not in archive_data:
        raise RuntimeError("Invalid DOCX package: missing word/document.xml")

    root = ET.fromstring(archive_data["word/document.xml"])
    paragraph_nodes = root.findall(".//w:p", ns)

    for para_id, anchor in para_anchor.items():
        if para_id not in issue_map:
            continue

        p_idx = anchor.get("paragraph_xml_index")
        if p_idx is None and anchor.get("kind") in {"paragraph", "word_main_paragraph"}:
            p_idx = anchor.get("paragraph_index")
        if p_idx is None:
            continue
        p_idx = int(p_idx)
        if not (0 <= p_idx < len(paragraph_nodes)):
            continue

        p_node = paragraph_nodes[p_idx]
        for issue in issue_map[para_id]:
            note = _build_comment_text(issue).replace("\n", " | ")
            r_node = ET.SubElement(p_node, f"{{{w_ns}}}r")
            t_node = ET.SubElement(r_node, f"{{{w_ns}}}t")
            t_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t_node.text = f" [{issue.issue_no}] REVIEW NOTE {note}"

    archive_data["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(output_path), "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, data in archive_data.items():
            zf.writestr(name, data)
    return output_path


def _convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    soffice = find_soffice()
    if not soffice:
        return None

    tmp_out = mkdtemp(prefix="doc_review_annot_pdf_")
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(tmp_out),
        str(docx_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        return None

    converted = tmp_out / f"{docx_path.stem}.pdf"
    if not converted.exists():
        return None

    final_path = output_dir / f"{docx_path.stem}.pdf"
    final_path.write_bytes(converted.read_bytes())
    return final_path


def _convert_docx_to_pdf_text_fallback(parse_result: ParseResult, issues: List[IssueItem], output_dir: Path) -> Path:
    out_pdf = output_dir / f"{parse_result.original_path.stem}.annotated.pdf"
    issue_map: Dict[str, List[IssueItem]] = defaultdict(list)
    for issue in issues:
        issue_map[issue.para_id].append(issue)

    doc = fitz.open()
    page = doc.new_page()
    margin = 48
    y = margin
    line_gap = 18

    for para in parse_result.paragraphs:
        para_start_y = y
        base_lines = _wrap_text_chars(para.text, 42)
        for line in base_lines:
            if y > page.rect.height - margin:
                page = doc.new_page()
                y = margin
                para_start_y = y
            page.insert_text((margin, y), line, fontname="helv", fontsize=11, color=(0, 0, 0))
            y += line_gap

        para_issues = issue_map.get(para.para_id, [])
        for issue in para_issues:
            label = f"[{issue.issue_no}] {issue.reason} | 建议: {issue.suggestion}"
            for line in _wrap_text_chars(label, 40):
                if y > page.rect.height - margin:
                    page = doc.new_page()
                    y = margin
                page.insert_text((margin + 12, y), line, fontname="helv", fontsize=10, color=(0.87, 0.2, 0.0))
                y += line_gap

            try:
                rect = fitz.Rect(margin, para_start_y - 12, page.rect.width - margin, min(y + 4, page.rect.height - margin))
                hl = page.add_highlight_annot(rect)
                hl.set_info(content=_build_comment_text(issue), title="DocReview")
                hl.update()
            except Exception:
                pass
        y += 6

    doc.save(str(out_pdf), garbage=4, deflate=True)
    doc.close()
    return out_pdf


def _wrap_text_chars(text: str, width: int) -> List[str]:
    if len(text) <= width:
        return [text]
    return [text[i : i + width] for i in range(0, len(text), width)]


def _render_pdf_source_to_markdown(parse_result: ParseResult) -> str:
    raw_md = str(parse_result.metadata.get("raw_markdown", "")).strip()
    if raw_md:
        return raw_md
    return "\n\n".join(p.text for p in parse_result.paragraphs if p.text.strip())


def apply_annotations(parse_result: ParseResult, issues: List[IssueItem], output_dir: str | Path) -> Tuple[Path, Optional[Path]]:
    output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    suffix = parse_result.normalized_path.suffix.lower()
    out_name = f"{parse_result.original_path.stem}.annotated{suffix if suffix in {'.docx', '.pdf'} else '.docx'}"
    output_path = output_dir / out_name

    if parse_result.source_type == "docx":
        annotated_docx = _apply_docx_annotations(parse_result, issues, output_path)
        return annotated_docx, None

    if parse_result.source_type == "pdf":
        source_md = _render_pdf_source_to_markdown(parse_result)
        annotated_md = build_annotated_markdown(parse_result, issues)
        if not annotated_md.strip():
            annotated_md = source_md

        md_path = output_dir / f"{parse_result.original_path.stem}.annotated.md"
        visual_images = []
        for item in parse_result.metadata.get("visual_region_images", []) or []:
            if isinstance(item, dict) and item.get("image_path"):
                visual_images.append(str(item["image_path"]))

        _, bundle_path = write_markdown_bundle(annotated_md, md_path, extra_files=visual_images)

        pdf_path = output_dir / f"{parse_result.original_path.stem}.annotated.pdf"
        issue_payload = [
            {
                "issue_id": issue.issue_id,
                "issue_no": issue.issue_no,
                "para_id": issue.para_id,
                "page_no": issue.page_no,
                "category": issue.category,
                "severity": issue.severity,
                "confidence": issue.confidence,
                "reason": issue.reason,
                "suggestion": issue.suggestion,
                "original_text": issue.original_text,
                "modified_text": issue.modified_text,
                "snippet": issue.snippet,
                "review_input_text": issue.review_input_text,
                "review_round": issue.review_round,
                "bbox_list": issue.bbox_list,
                "anchor": issue.anchor,
            }
            for issue in issues
        ]
        annotate_pdf_from_issues(
            input_pdf=parse_result.normalized_path,
            issues=issue_payload,
            output_pdf=pdf_path,
        )
        return pdf_path, bundle_path

    raise ValueError(f"Unsupported source type: {parse_result.source_type}")
