from __future__ import annotations

import hashlib
import re
import subprocess
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List

import fitz  # PyMuPDF

try:
    from ..models import ParagraphUnit, ParseResult, PdfParseBackend
    from .doc2x import parse_pdf_to_doc2x
    from .office import find_soffice
    from .paddle_ocr import parse_pdf_to_markdown
    from .temp_paths import mkdtemp
except ImportError:  # pragma: no cover
    from models import ParagraphUnit, ParseResult, PdfParseBackend
    from services.doc2x import parse_pdf_to_doc2x
    from services.office import find_soffice
    from services.paddle_ocr import parse_pdf_to_markdown
    from services.temp_paths import mkdtemp


def _normalize_text(text: str) -> str:
    return "\n".join([line.strip() for line in text.splitlines() if line.strip()]).strip()


def _extract_docx_paragraph_meta(paragraph) -> Dict[str, object]:
    align = None
    space_before = None
    space_after = None
    line_spacing = None
    style_name = None
    try:
        fmt = paragraph.paragraph_format
        align = paragraph.alignment
        space_before = fmt.space_before.pt if fmt.space_before is not None else None
        space_after = fmt.space_after.pt if fmt.space_after is not None else None
        line_spacing = fmt.line_spacing
        style_name = getattr(getattr(paragraph, "style", None), "name", None)
    except Exception:
        pass
    return {
        "alignment": str(align) if align is not None else None,
        "space_before": space_before,
        "space_after": space_after,
        "line_spacing": line_spacing,
        "style_name": style_name,
        "raw_text": getattr(paragraph, "text", ""),
    }


def _extract_docx_structure_fields(paragraph) -> Dict[str, object]:
    meta = _extract_docx_paragraph_meta(paragraph)
    raw_runs: List[str] = []
    hard_breaks = 0
    try:
        for run in paragraph.runs:
            run_text = getattr(run, "text", "")
            if run_text:
                raw_runs.append(run_text)
            hard_breaks += int(run_text.count("\n") + run_text.count("\v"))
    except Exception:
        pass
    meta["raw_runs_text"] = raw_runs
    meta["hard_break_count"] = hard_breaks
    try:
        formula_texts = _extract_omml_texts(paragraph._element)
    except Exception:
        formula_texts = []
    meta["formula_texts"] = formula_texts
    return meta


def _extract_omml_texts(node: ET.Element) -> List[str]:
    texts: List[str] = []

    def walk(current: ET.Element) -> None:
        local = _xml_local_name(current.tag)
        if local in {"oMath", "oMathPara"}:
            formula = _omml_to_text(current).strip()
            if formula:
                texts.append(formula)
            return
        for child in list(current):
            walk(child)

    walk(node)
    return texts


def _xml_local_name(tag: object) -> str:
    return str(tag).rsplit("}", 1)[-1]


def _omml_group_text(node: ET.Element, group_name: str) -> str:
    for child in list(node):
        if _xml_local_name(child.tag) == group_name:
            return _omml_to_text(child)
    return ""


def _omml_power_text(text: str) -> str:
    text = text.strip()
    if not text:
        return ""
    return text if len(text) == 1 else "{" + text + "}"


def _omml_to_text(node: ET.Element) -> str:
    local = _xml_local_name(node.tag)
    if local == "t":
        return node.text or ""
    if local == "sSup":
        base = _omml_group_text(node, "e")
        sup = _omml_group_text(node, "sup")
        return f"{base}^{_omml_power_text(sup)}" if sup else base
    if local == "sSub":
        base = _omml_group_text(node, "e")
        sub = _omml_group_text(node, "sub")
        return f"{base}_{_omml_power_text(sub)}" if sub else base
    if local == "sSubSup":
        base = _omml_group_text(node, "e")
        sub = _omml_group_text(node, "sub")
        sup = _omml_group_text(node, "sup")
        if sub and sup:
            return f"{base}_{_omml_power_text(sub)}^{_omml_power_text(sup)}"
        if sub:
            return f"{base}_{_omml_power_text(sub)}"
        if sup:
            return f"{base}^{_omml_power_text(sup)}"
        return base
    if local == "f":
        num = _omml_group_text(node, "num")
        den = _omml_group_text(node, "den")
        return f"({num})/({den})" if num or den else ""
    if local == "rad":
        deg = _omml_group_text(node, "deg")
        base = _omml_group_text(node, "e")
        return f"root[{deg}]({base})" if deg else f"sqrt({base})"
    if local == "nary":
        operator = _omml_group_text(node, "chr")
        sub = _omml_group_text(node, "sub")
        sup = _omml_group_text(node, "sup")
        expr = _omml_group_text(node, "e")
        prefix = operator or "∑"
        if sub:
            prefix += f"_{_omml_power_text(sub)}"
        if sup:
            prefix += f"^{_omml_power_text(sup)}"
        return prefix + expr

    parts = [_omml_to_text(child) for child in list(node)]
    formula = "".join(parts).strip()
    if formula:
        return formula
    if node.text:
        return node.text
    return ""


def _append_formula_text(text: str, formula_texts: List[str]) -> str:
    formulas = [x.strip() for x in formula_texts if x and x.strip()]
    if not formulas:
        return text
    suffix = " ".join(f"[公式: {x}]" for x in formulas)
    return _normalize_text(f"{text}\n{suffix}")


def _build_doc_id(path: Path) -> str:
    payload = f"{path.resolve()}::{path.stat().st_mtime_ns}::{path.stat().st_size}"
    return hashlib.sha1(payload.encode("utf-8")).hexdigest()[:16]


def _convert_doc_to_docx_if_needed(path: Path) -> Path:
    if path.suffix.lower() != ".doc":
        return path

    soffice = find_soffice()
    if not soffice:
        raise RuntimeError(".doc files require LibreOffice (soffice) for conversion, but it is not installed.")

    out_dir = mkdtemp(prefix="doc_review_doc_convert_")
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(out_dir),
        str(path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"Failed to convert .doc to .docx: {result.stderr.strip()}")

    converted = out_dir / f"{path.stem}.docx"
    if not converted.exists():
        raise RuntimeError("Conversion finished but output .docx was not found.")
    return converted


def _parse_docx(path: Path, normalized_path: Path, doc_id: str) -> ParseResult:
    try:
        from docx import Document
    except Exception:
        return _parse_docx_openxml(path, normalized_path, doc_id)

    doc = Document(str(normalized_path))
    paragraphs: List[ParagraphUnit] = []

    para_counter = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        meta = _extract_docx_structure_fields(paragraph)
        text = _normalize_text(str(meta.get("raw_text") or getattr(paragraph, "text", "")))
        text = _append_formula_text(text, meta.get("formula_texts", []))
        if not text:
            continue
        para_counter += 1
        paragraphs.append(
            ParagraphUnit(
                doc_id=doc_id,
                para_id=f"para-{para_counter:05d}",
                page_no=None,
                text=text,
                source_type="docx",
                anchor={
                    "kind": "paragraph",
                    "paragraph_index": idx,
                    "docx_meta": meta,
                    "raw_runs_text": meta.get("raw_runs_text", []),
                    "hard_break_count": meta.get("hard_break_count", 0),
                    "alignment": meta.get("alignment"),
                    "space_before": meta.get("space_before"),
                    "space_after": meta.get("space_after"),
                    "line_spacing": meta.get("line_spacing"),
                    "style_name": meta.get("style_name"),
                    "raw_text": meta.get("raw_text", ""),
                },
            )
        )

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    meta = _extract_docx_structure_fields(paragraph)
                    text = _normalize_text(str(meta.get("raw_text") or getattr(paragraph, "text", "")))
                    text = _append_formula_text(text, meta.get("formula_texts", []))
                    if not text:
                        continue
                    para_counter += 1
                    paragraphs.append(
                        ParagraphUnit(
                            doc_id=doc_id,
                            para_id=f"para-{para_counter:05d}",
                            page_no=None,
                            text=text,
                            source_type="docx",
                            anchor={
                                "kind": "table_cell",
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                                "paragraph_index": p_idx,
                                "docx_meta": meta,
                            },
                        )
                    )

    return ParseResult(
        source_type="docx",
        original_path=path,
        normalized_path=normalized_path,
        doc_id=doc_id,
        paragraphs=paragraphs,
        metadata={
            "paragraph_count": len(paragraphs),
            "docx_parser": "python-docx",
            "docx_structure": True,
        },
    )


def _parse_docx_openxml(path: Path, normalized_path: Path, doc_id: str) -> ParseResult:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}

    with zipfile.ZipFile(str(normalized_path), "r") as zf:
        try:
            xml_bytes = zf.read("word/document.xml")
        except KeyError as exc:
            raise RuntimeError("Invalid DOCX package: missing word/document.xml") from exc

    root = ET.fromstring(xml_bytes)
    paragraph_nodes = root.findall(".//w:p", ns)
    paragraphs: List[ParagraphUnit] = []
    para_counter = 0

    for p_idx, p_node in enumerate(paragraph_nodes):
        texts = [t.text or "" for t in p_node.findall(".//w:t", ns)]
        formula_texts = _extract_omml_texts(p_node)
        text = _normalize_text("".join(texts))
        text = _append_formula_text(text, formula_texts)
        if not text:
            continue

        raw_text, hard_breaks = _extract_openxml_paragraph_raw_text(p_node, ns)
        raw_text = _append_formula_text(_normalize_text(raw_text), formula_texts)
        ppr = p_node.find("./w:pPr", ns)
        align = None
        space_before = None
        space_after = None
        line_spacing = None
        if ppr is not None:
            jc = ppr.find("./w:jc", ns)
            if jc is not None:
                align = jc.attrib.get(f"{{{w_ns}}}val")
            spacing = ppr.find("./w:spacing", ns)
            if spacing is not None:
                space_before = spacing.attrib.get(f"{{{w_ns}}}before")
                space_after = spacing.attrib.get(f"{{{w_ns}}}after")
                line_spacing = spacing.attrib.get(f"{{{w_ns}}}line")

        para_counter += 1
        paragraphs.append(
            ParagraphUnit(
                doc_id=doc_id,
                para_id=f"para-{para_counter:05d}",
                page_no=None,
                text=text,
                source_type="docx",
                anchor={
                    "kind": "paragraph_xml",
                    "paragraph_xml_index": p_idx,
                    "raw_runs_text": [t.text or "" for t in p_node.findall(".//w:t", ns) if (t.text or "")],
                    "docx_meta": {
                        "alignment": align,
                        "space_before": space_before,
                        "space_after": space_after,
                        "line_spacing": line_spacing,
                        "style_name": None,
                        "raw_text": raw_text,
                        "formula_texts": formula_texts,
                    },
                    "hard_break_count": hard_breaks,
                    "raw_text": raw_text,
                },
            )
        )

    return ParseResult(
        source_type="docx",
        original_path=path,
        normalized_path=normalized_path,
        doc_id=doc_id,
        paragraphs=paragraphs,
        metadata={"paragraph_count": len(paragraphs), "docx_parser": "openxml_fallback", "docx_structure": True},
    )


def _extract_openxml_paragraph_raw_text(p_node: ET.Element, ns: Dict[str, str]) -> tuple[str, int]:
    w_ns = ns["w"]
    parts: List[str] = []
    hard_breaks = 0
    for node in p_node.iter():
        if node.tag == f"{{{w_ns}}}t":
            parts.append(node.text or "")
        elif node.tag in {f"{{{w_ns}}}br", f"{{{w_ns}}}cr"}:
            parts.append("\n")
            hard_breaks += 1
        elif node.tag == f"{{{w_ns}}}tab":
            parts.append("\t")
    return "".join(parts), hard_breaks


def _parse_docx_word_helper(path: Path, normalized_path: Path, doc_id: str) -> ParseResult:
    try:
        from docx import Document
    except Exception:
        parsed = _parse_docx_openxml(path, normalized_path, doc_id)
        try:
            with zipfile.ZipFile(str(normalized_path), "r") as zf:
                root = ET.fromstring(zf.read("word/document.xml"))
            paragraph_nodes = root.findall(".//w:p", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        except Exception:
            paragraph_nodes = []
        for idx, para in enumerate(parsed.paragraphs):
            raw_text = para.text
            hard_breaks = 0
            if idx < len(paragraph_nodes):
                raw_text, hard_breaks = _extract_openxml_paragraph_raw_text(
                    paragraph_nodes[idx],
                    {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
                )
            para.anchor["kind"] = "word_helper_paragraph"
            para.anchor["docx_meta"] = {"raw_text": raw_text}
            para.anchor["hard_break_count"] = hard_breaks
            if raw_text:
                para.text = _normalize_text(raw_text)
            para.anchor["alignment"] = None
        parsed.metadata["word_helper_source"] = str(path)
        parsed.metadata["word_helper_raw"] = True
        return parsed

    doc = Document(str(normalized_path))
    paragraphs: List[ParagraphUnit] = []
    para_counter = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        raw_text = getattr(paragraph, "text", "")
        text = _normalize_text(raw_text)
        structure_meta = _extract_docx_structure_fields(paragraph)
        text = _append_formula_text(text, structure_meta.get("formula_texts", []))
        if not text:
            continue
        para_counter += 1
        meta = _extract_docx_paragraph_meta(paragraph)
        meta["formula_texts"] = structure_meta.get("formula_texts", [])
        paragraphs.append(
            ParagraphUnit(
                doc_id=doc_id,
                para_id=f"para-{para_counter:05d}",
                page_no=None,
                text=text,
                source_type="docx",
                anchor={
                    "kind": "word_helper_paragraph",
                    "paragraph_index": idx,
                    "docx_meta": meta,
                    "raw_runs_text": structure_meta["raw_runs_text"],
                    "hard_break_count": structure_meta["hard_break_count"],
                    "alignment": meta.get("alignment"),
                },
            )
        )

    return ParseResult(
        source_type="docx",
        original_path=path,
        normalized_path=normalized_path,
        doc_id=doc_id,
        paragraphs=paragraphs,
        metadata={
            "paragraph_count": len(paragraphs),
            "docx_parser": "python-docx",
            "word_helper_source": str(path),
            "word_helper_raw": True,
            "docx_structure": True,
        },
    )


def _parse_docx_structure(path: Path, normalized_path: Path, doc_id: str) -> ParseResult:
    try:
        from docx import Document
    except Exception:
        return _parse_docx_openxml(path, normalized_path, doc_id)

    doc = Document(str(normalized_path))
    paragraphs: List[ParagraphUnit] = []
    para_counter = 0

    for idx, paragraph in enumerate(doc.paragraphs):
        meta = _extract_docx_structure_fields(paragraph)
        text = _normalize_text(str(meta.get("raw_text") or getattr(paragraph, "text", "")))
        text = _append_formula_text(text, meta.get("formula_texts", []))
        if not text:
            continue
        para_counter += 1
        paragraphs.append(
            ParagraphUnit(
                doc_id=doc_id,
                para_id=f"para-{para_counter:05d}",
                page_no=None,
                text=text,
                source_type="docx",
                anchor={
                    "kind": "word_main_paragraph",
                    "paragraph_index": idx,
                    "docx_meta": meta,
                    "raw_runs_text": meta.get("raw_runs_text", []),
                    "hard_break_count": meta.get("hard_break_count", 0),
                    "alignment": meta.get("alignment"),
                    "space_before": meta.get("space_before"),
                    "space_after": meta.get("space_after"),
                    "line_spacing": meta.get("line_spacing"),
                    "style_name": meta.get("style_name"),
                    "raw_text": meta.get("raw_text", ""),
                },
            )
        )

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    meta = _extract_docx_structure_fields(paragraph)
                    text = _normalize_text(str(meta.get("raw_text") or getattr(paragraph, "text", "")))
                    text = _append_formula_text(text, meta.get("formula_texts", []))
                    if not text:
                        continue
                    para_counter += 1
                    paragraphs.append(
                        ParagraphUnit(
                            doc_id=doc_id,
                            para_id=f"para-{para_counter:05d}",
                            page_no=None,
                            text=text,
                            source_type="docx",
                            anchor={
                                "kind": "word_main_table_cell",
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                                "paragraph_index": p_idx,
                                "docx_meta": meta,
                                "raw_runs_text": meta.get("raw_runs_text", []),
                                "hard_break_count": meta.get("hard_break_count", 0),
                                "alignment": meta.get("alignment"),
                                "space_before": meta.get("space_before"),
                                "space_after": meta.get("space_after"),
                                "line_spacing": meta.get("line_spacing"),
                                "style_name": meta.get("style_name"),
                                "raw_text": meta.get("raw_text", ""),
                            },
                        )
                    )

    return ParseResult(
        source_type="docx",
        original_path=path,
        normalized_path=normalized_path,
        doc_id=doc_id,
        paragraphs=paragraphs,
        metadata={
            "paragraph_count": len(paragraphs),
            "docx_parser": "python-docx",
            "docx_structure": True,
        },
    )


def _is_markdown_image_block(text: str) -> bool:
    s = text.strip()
    return bool(re.match(r"^!\[[^\]]*\]\([^)]+\)\s*$", s))


def _split_markdown_paragraphs(markdown_text: str, include_images: bool = False) -> List[str]:
    if not markdown_text:
        return []

    cleaned_lines = []
    for line in markdown_text.splitlines():
        s = line.strip()
        if not s:
            cleaned_lines.append("")
            continue
        if not include_images and s.startswith("![") and "](" in s:
            continue
        cleaned_lines.append(line)

    blocks = re.split(r"\n\s*\n+", "\n".join(cleaned_lines))
    return [_normalize_text(b) for b in blocks if _normalize_text(b)]


def _parse_pdf(path: Path, doc_id: str) -> ParseResult:
    backend = "paddle_api"
    paddle_result = parse_pdf_to_markdown(path)
    page_paragraphs: List[tuple[int | None, str]] = []
    for page in getattr(paddle_result, "pages", []) or []:
        page_no = int(getattr(page, "page_no", 0) or 0) or None
        for chunk in _split_markdown_paragraphs(getattr(page, "markdown_text", ""), include_images=True):
            page_paragraphs.append((page_no, chunk))
    if not page_paragraphs:
        page_paragraphs = [
            (None, chunk)
            for chunk in _split_markdown_paragraphs(paddle_result.markdown_text, include_images=True)
        ]
    if not page_paragraphs:
        raise RuntimeError("Paddle OCR produced no usable markdown paragraphs.")

    paragraphs: List[ParagraphUnit] = []
    for idx, (page_no, chunk) in enumerate(page_paragraphs, start=1):
        paragraphs.append(
            ParagraphUnit(
                doc_id=doc_id,
                para_id=f"para-{idx:05d}",
                page_no=page_no,
                text=chunk,
                bbox_list=[],
                source_type="pdf",
                anchor={
                    "kind": "pdf_paddle_api_image" if _is_markdown_image_block(chunk) else "pdf_paddle_api_markdown",
                    "paragraph_index": idx - 1,
                    "page_no": page_no,
                    "reviewable": not _is_markdown_image_block(chunk),
                },
            )
        )

    page_count = 0
    try:
        with fitz.open(str(path)) as pdf_doc:
            page_count = pdf_doc.page_count
    except Exception:
        page_count = len({p.page_no for p in paddle_result.pages if p.page_no})

    metadata: Dict[str, object] = {
        "page_count": page_count,
        "ocr_used": True,
        "pdf_parser_backend": backend,
        "ocr_backend": "paddle_api",
        "markdown_used": True,
        "raw_markdown": paddle_result.markdown_text,
        "paddle_job_meta": paddle_result.meta,
    }

    return ParseResult(
        source_type="pdf",
        original_path=path,
        normalized_path=path,
        doc_id=doc_id,
        paragraphs=paragraphs,
        metadata=metadata,
    )


def _parse_pdf_doc2x(path: Path, doc_id: str) -> ParseResult:
    backend = "doc2x_v3_json"
    doc2x_result = parse_pdf_to_doc2x(path)
    page_paragraphs: List[tuple[int | None, str]] = []
    for page in getattr(doc2x_result, "pages", []) or []:
        page_no = int(getattr(page, "page_no", 0) or 0) or None
        for chunk in _split_markdown_paragraphs(getattr(page, "markdown_text", ""), include_images=True):
            page_paragraphs.append((page_no, chunk))
    if not page_paragraphs:
        page_paragraphs = [
            (None, chunk)
            for chunk in _split_markdown_paragraphs(doc2x_result.markdown_text, include_images=True)
        ]
    if not page_paragraphs:
        raise RuntimeError("Doc2X produced no usable markdown paragraphs.")

    paragraphs: List[ParagraphUnit] = []
    for idx, (page_no, chunk) in enumerate(page_paragraphs, start=1):
        paragraphs.append(
            ParagraphUnit(
                doc_id=doc_id,
                para_id=f"para-{idx:05d}",
                page_no=page_no,
                text=chunk,
                bbox_list=[],
                source_type="pdf",
                anchor={
                    "kind": "pdf_doc2x_image" if _is_markdown_image_block(chunk) else "pdf_doc2x_markdown",
                    "paragraph_index": idx - 1,
                    "page_no": page_no,
                    "reviewable": not _is_markdown_image_block(chunk),
                },
            )
        )

    page_count = 0
    try:
        with fitz.open(str(path)) as pdf_doc:
            page_count = pdf_doc.page_count
    except Exception:
        page_count = int(doc2x_result.meta.get("page_count") or 0)

    metadata: Dict[str, object] = {
        "page_count": page_count,
        "ocr_used": True,
        "pdf_parser_backend": backend,
        "ocr_backend": "doc2x",
        "markdown_used": True,
        "raw_markdown": doc2x_result.markdown_text,
        "doc2x_job_meta": doc2x_result.meta,
        "doc2x_raw_result": doc2x_result.raw_result,
        "doc2x_layout_regions": doc2x_result.layout_regions,
    }

    return ParseResult(
        source_type="pdf",
        original_path=path,
        normalized_path=path,
        doc_id=doc_id,
        paragraphs=paragraphs,
        metadata=metadata,
    )


def parse_document(file_path: str | Path, pdf_parse_backend: PdfParseBackend = "paddle") -> ParseResult:
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in {".doc", ".docx", ".pdf"}:
        raise ValueError("Unsupported file type. Only .doc/.docx/.pdf are allowed.")

    normalized = _convert_doc_to_docx_if_needed(path)
    doc_id = _build_doc_id(normalized)
    backend = pdf_parse_backend if pdf_parse_backend in {"paddle", "doc2x"} else "paddle"

    if normalized.suffix.lower() == ".docx":
        if path.suffix.lower() == ".pdf":
            return _parse_pdf_doc2x(path, doc_id) if backend == "doc2x" else _parse_pdf(path, doc_id)
        return _parse_docx_structure(path, normalized, doc_id)
    return _parse_pdf_doc2x(path, doc_id) if backend == "doc2x" else _parse_pdf(path, doc_id)


def parse_word_helper_document(file_path: str | Path) -> ParseResult:
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    suffix = path.suffix.lower()
    if suffix not in {".doc", ".docx"}:
        raise ValueError("Unsupported file type. Only .doc/.docx are allowed for Word helper parsing.")
    normalized = _convert_doc_to_docx_if_needed(path)
    doc_id = _build_doc_id(normalized)
    return _parse_docx_word_helper(path, normalized, doc_id)
