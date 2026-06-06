from __future__ import annotations

import html
from pathlib import Path
from typing import List
import zipfile
import xml.etree.ElementTree as ET
import subprocess
import os
import fitz  # PyMuPDF

try:
    from .office import find_soffice
except ImportError:  # pragma: no cover
    from services.office import find_soffice

PROJECT_ROOT = Path(__file__).resolve().parents[1]
PREVIEW_DIR = PROJECT_ROOT / "outputs" / "previews"


def _docx_to_html(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return _docx_to_html_openxml(path)

    doc = Document(str(path))
    blocks: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(f"<p style='color:#111827; margin:0 0 10px 0;'>{html.escape(text)}</p>")

    for table in doc.tables:
        rows_html = []
        for row in table.rows:
            cells = "".join(
                [
                    f"<td style='color:#111827; vertical-align:top;'>{html.escape(cell.text.strip())}</td>"
                    for cell in row.cells
                ]
            )
            rows_html.append(f"<tr>{cells}</tr>")
        blocks.append("<table border='1' cellspacing='0' cellpadding='4'>" + "".join(rows_html) + "</table>")

    body = "\n".join(blocks) if blocks else "<p><em>No text content found.</em></p>"
    return (
        "<div style='font-family: sans-serif; color:#111827; line-height:1.7; "
        "padding: 12px; max-height: 72vh; overflow: auto; background: #ffffff;'>"
        + body
        + "</div>"
    )


def _docx_to_html_openxml(path: Path) -> str:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    blocks: List[str] = []

    with zipfile.ZipFile(str(path), "r") as zf:
        xml_bytes = zf.read("word/document.xml")
    root = ET.fromstring(xml_bytes)

    for p_node in root.findall(".//w:p", ns):
        text = "".join((t.text or "") for t in p_node.findall(".//w:t", ns)).strip()
        if text:
            blocks.append(f"<p style='color:#111827; margin:0 0 10px 0;'>{html.escape(text)}</p>")

    body = "\n".join(blocks) if blocks else "<p><em>No text content found.</em></p>"
    return (
        "<div style='font-family: sans-serif; color:#111827; line-height:1.7; "
        "padding: 12px; max-height: 72vh; overflow: auto; background: #ffffff;'>"
        + body
        + "</div>"
    )


def build_preview_payload(file_path: str | Path):
    path = Path(file_path).resolve()
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return {
            "type": "pdf",
            "pdf_path": str(path),
            "html": None,
        }

    if suffix in {".docx", ".doc"}:
        preview_pdf_path, preview_error = _convert_office_to_pdf(path)
        if preview_pdf_path is not None:
            return {
                "type": "pdf",
                "pdf_path": str(preview_pdf_path),
                "html": None,
                "preview_origin": "converted",
                "preview_error": None,
            }

        text_pdf_path = _convert_docx_to_text_pdf(path)
        if text_pdf_path is not None:
            return {
                "type": "pdf",
                "pdf_path": str(text_pdf_path),
                "html": None,
                "preview_origin": "converted_text_pdf",
                "preview_error": preview_error,
            }

        html_payload = _docx_to_html(path)
        return {
            "type": "docx",
            "pdf_path": None,
            "html": html_payload,
            "preview_origin": "html_fallback",
            "preview_error": preview_error,
        }

    raise ValueError("Unsupported file type for preview")


def _convert_office_to_pdf(path: Path) -> tuple[Path | None, str | None]:
    soffice = find_soffice()
    if not soffice:
        return None, "soffice_not_found"

    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    out_dir = PREVIEW_DIR / f"{path.stem}_{path.stat().st_mtime_ns}"
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(path),
    ]
    try:
        timeout = max(1, int(os.getenv("DOC_REVIEW_PREVIEW_CONVERT_TIMEOUT", "15")))
    except ValueError:
        timeout = 15
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
    except subprocess.TimeoutExpired:
        return None, f"convert_timeout_after_{timeout}s"
    if result.returncode != 0:
        msg = result.stderr.strip() or result.stdout.strip() or "convert_failed"
        return None, msg

    out_pdf = out_dir / f"{path.stem}.pdf"
    if not out_pdf.exists():
        return None, "converted_pdf_missing"
    return out_pdf.resolve(), None


def _convert_docx_to_text_pdf(path: Path) -> Path | None:
    try:
        paragraphs = _extract_docx_text_lines(path)
    except Exception:
        return None
    if not paragraphs:
        return None

    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    out_dir = PREVIEW_DIR / f"{path.stem}_{path.stat().st_mtime_ns}_textpdf"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_pdf = out_dir / f"{path.stem}.pdf"

    doc = fitz.open()
    page = doc.new_page()
    margin = 48
    y = margin
    line_gap = 18

    for para in paragraphs:
        chunks = _wrap_text_chars(para, 42)
        for line in chunks:
            if y > page.rect.height - margin:
                page = doc.new_page()
                y = margin
            page.insert_text((margin, y), line, fontname="helv", fontsize=11, color=(0, 0, 0))
            y += line_gap
        y += 4

    doc.save(str(out_pdf), garbage=4, deflate=True)
    doc.close()
    return out_pdf.resolve()


def _extract_docx_text_lines(path: Path) -> List[str]:
    try:
        from docx import Document

        doc = Document(str(path))
        lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        return lines
    except Exception:
        return _extract_docx_text_lines_openxml(path)


def _extract_docx_text_lines_openxml(path: Path) -> List[str]:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    lines: List[str] = []
    with zipfile.ZipFile(str(path), "r") as zf:
        xml_bytes = zf.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    for p_node in root.findall(".//w:p", ns):
        text = "".join((t.text or "") for t in p_node.findall(".//w:t", ns)).strip()
        if text:
            lines.append(text)
    return lines


def _wrap_text_chars(text: str, width: int) -> List[str]:
    if len(text) <= width:
        return [text]
    return [text[i : i + width] for i in range(0, len(text), width)]
